import os
from typing import List, Dict, Any
from langchain_core.documents import Document

def load_document(file_path: str, metadata: Dict[str, Any] = None) -> List[Document]:
    """
    Loads a PDF, DOCX, TXT, or MD document from file_path and attaches metadata.
    """
    if metadata is None:
        metadata = {}

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()
    docs = []

    base_meta = {
        "source": os.path.basename(file_path),
        "file_path": file_path,
        **metadata
    }

    try:
        if ext in [".txt", ".md"]:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            docs.append(Document(page_content=content, metadata=base_meta))

        elif ext == ".pdf":
            try:
                from pypdf import PdfReader
                reader = PdfReader(file_path)
                for i, page in enumerate(reader.pages):
                    text = page.extract_text() or ""
                    if text.strip():
                        page_meta = {**base_meta, "page": i + 1}
                        docs.append(Document(page_content=text, metadata=page_meta))
            except Exception as e:
                print(f"[Loader] PDF parsing fallback for {file_path}: {e}")
                from langchain_community.document_loaders import PyPDFLoader
                loader = PyPDFLoader(file_path)
                loaded_docs = loader.load()
                for doc in loaded_docs:
                    doc.metadata.update(base_meta)
                docs.extend(loaded_docs)

        elif ext == ".docx":
            import docx
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            text = "\n\n".join(full_text)
            docs.append(Document(page_content=text, metadata=base_meta))

        else:
            print(f"[Loader] Unsupported file extension {ext} for file {file_path}")

    except Exception as err:
        print(f"[Loader] Error loading {file_path}: {err}")

    return docs
